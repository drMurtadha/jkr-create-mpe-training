from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path("MPE_Hub_UAT_and_Enhancement_Documentation_original.docx")
OUTPUT = Path("MPE_Hub_UAT_and_Enhancement_Documentation_Improved.docx")


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def move_after(element, anchor):
    anchor.addnext(element)
    return element


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.extend([color, underline])
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.extend([run_pr, text_el])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def insert_revision_history(doc, after_table):
    heading = doc.add_paragraph("Revision history", style="Heading 2")
    table = doc.add_table(rows=3, cols=5)
    table.style = "Table Grid"
    headers = ["Version", "Date", "Prepared by", "Summary of change", "Approval status"]
    for i, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], value, bold=True)
        shade_cell(table.rows[0].cells[i], "D9EAF7")
    rows = [
        ["1.0", "18 July 2026", "Document author", "Initial controlled UAT baseline", "Superseded"],
        ["1.1", "18 July 2026", "Document author", "Clarified governance, execution status, evidence and sign-off requirements", "For review"],
    ]
    for row, values in zip(table.rows[1:], rows):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)
    set_repeat_table_header(table.rows[0])
    anchor = after_table._tbl
    move_after(heading._p, anchor)
    move_after(table._tbl, heading._p)


def insert_status_legend(doc, anchor_paragraph):
    heading = doc.add_paragraph("5.1 Test status definitions", style="Heading 2")
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    entries = [
        ("Status", "Definition"),
        ("Not Run", "The test has not started or required evidence is unavailable."),
        ("Pass", "The expected visible behavior and required downstream evidence are both confirmed."),
        ("Fail", "The actual result differs from the expected result; a defect ID is recorded."),
        ("Blocked", "The test cannot proceed because a prerequisite, environment or dependency is unavailable."),
    ]
    for row, values in zip(table.rows, entries):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value, bold=row is table.rows[0])
            if row is table.rows[0]:
                shade_cell(cell, "D9EAF7")
    set_repeat_table_header(table.rows[0])
    move_after(heading._p, anchor_paragraph._p)
    move_after(table._tbl, heading._p)


def insert_signoff(doc, anchor_paragraph):
    heading = doc.add_paragraph("14.1 Release decision and sign-off", style="Heading 2")
    intro = doc.add_paragraph(
        "Record the final decision only after all release gates in Section 10 have supporting evidence. "
        "Any conditional approval must identify residual risks, an accountable owner and a due date."
    )
    table = doc.add_table(rows=5, cols=5)
    table.style = "Table Grid"
    entries = [
        ["Role", "Name", "Decision", "Signature / confirmation", "Date"],
        ["System owner", "", "Approve / Conditional / Reject", "", ""],
        ["Technical owner", "", "Approve / Conditional / Reject", "", ""],
        ["UAT representative", "", "Approve / Conditional / Reject", "", ""],
        ["Training facilitator", "", "Approve / Conditional / Reject", "", ""],
    ]
    for ri, (row, values) in enumerate(zip(table.rows, entries)):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value, bold=ri == 0)
            if ri == 0:
                shade_cell(cell, "D9EAF7")
    set_repeat_table_header(table.rows[0])
    move_after(heading._p, anchor_paragraph._p)
    move_after(intro._p, heading._p)
    move_after(table._tbl, intro._p)


doc = Document(SOURCE)

# Precise editorial improvements.
replacements = {
    "Fictional-user test pack, UAT scenarios, defect register and improvement roadmap":
        "Controlled UAT pack, synthetic test data, defect register and improvement roadmap",
    "This document converts the initial usability review into a repeatable UAT pack. It defines a fictional participant, safe test records, end-to-end scenarios, evidence requirements, observed issues, recommended enhancements and release gates. It is designed to support both immediate training validation and later product hardening.":
        "This document turns the initial usability review into a repeatable, evidence-based UAT pack. It defines a fictional participant, controlled synthetic records, end-to-end scenarios, evidence requirements, observed issues, prioritized enhancements and release gates. It supports immediate training validation and provides a governed path toward product hardening; it does not constitute production approval.",
    "2.2 Out of scope for the initial observation": "2.2 Constraints from the initial observation",
    "Record the actual result, evidence link or screenshot filename, tester, date and defect ID for every test. “Pass” requires both the visible success state and the downstream storage evidence where applicable.":
        "For every test, record the actual result, evidence link or screenshot filename, tester, execution date, build/version and defect ID. A test is “Pass” only when the visible success state and all applicable downstream storage evidence are confirmed.",
}
for paragraph in doc.paragraphs:
    if paragraph.text in replacements:
        paragraph.text = replacements[paragraph.text]

# Document-control metadata.
control = doc.tables[0]
for row in control.rows:
    key = row.cells[0].text.strip()
    if key == "Document status":
        set_cell_text(row.cells[1], "Version 1.1 — controlled baseline for review and execution")
    elif key == "Version":
        set_cell_text(row.cells[1], "1.1")
new_row = control.add_row()
set_cell_text(new_row.cells[0], "Document owner")
set_cell_text(new_row.cells[1], "To be assigned by the system owner before execution")
new_row = control.add_row()
set_cell_text(new_row.cells[0], "Review cycle")
set_cell_text(new_row.cells[1], "At each release candidate and after any material workflow, schema or security change")
insert_revision_history(doc, control)

# Repeat table headings and prevent row splitting for cleaner pagination.
for table in doc.tables:
    if table.rows:
        set_repeat_table_header(table.rows[0])
    for row in table.rows:
        keep_row_together(row)

# Clarify the compact priority heading without relying on table position.
for table in doc.tables:
    if table.rows and any(cell.text.strip() == "Pri." for cell in table.rows[0].cells):
        for cell in table.rows[0].cells:
            if cell.text.strip() == "Pri.":
                set_cell_text(cell, "Priority", bold=True)

# Add execution definitions immediately after the Section 5 instruction.
section5_instruction = next(p for p in doc.paragraphs if p.text.startswith("For every test, record"))
insert_status_legend(doc, section5_instruction)

# Add formal approval record after the final handover checklist item.
last_check = next(p for p in doc.paragraphs if "Release decision and residual risks signed" in p.text)
insert_signoff(doc, last_check)

# Convert appendix URLs to live hyperlinks while retaining visible destinations.
for paragraph in doc.paragraphs:
    if ": https://" in paragraph.text and (
        paragraph.text.startswith("MPE Hub:") or paragraph.text.startswith("Module 5")
    ):
        label, url = paragraph.text.split(": ", 1)
        paragraph.clear()
        lead = paragraph.add_run(label + ": ")
        lead.bold = True
        add_hyperlink(paragraph, url, url)

# Core metadata and accessible table descriptions where Word supports them.
doc.core_properties.title = "MPE Hub and Module 5 — UAT and Enhancement Documentation"
doc.core_properties.subject = "Controlled user acceptance testing, evidence, defects and release governance"
doc.core_properties.keywords = "MPE Hub, Module 5, UAT, synthetic data, accessibility, release gates"
doc.core_properties.comments = "Improved controlled baseline, version 1.1"

# Normalize body typography without disturbing display styles.
styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10)
for style_name in ["Heading 1", "Heading 2", "Title", "Subtitle"]:
    if style_name in styles:
        styles[style_name].font.name = "Arial"

doc.save(OUTPUT)
print(OUTPUT.resolve())
